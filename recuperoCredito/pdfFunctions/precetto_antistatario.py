# import textwrap
# from reportlab.lib.pagesizes import letter
# from reportlab.lib.units import inch
# from reportlab.pdfgen import canvas
# from django.http import HttpResponse

import io
from django.http import FileResponse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def precetto_antistatario(model=None):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(4.75)
        section.bottom_margin = Cm(3.25)
        section.left_margin = Cm(2.75)
        section.right_margin = Cm(3.75)

    doc_style = document.styles

    para_upper_left_style = doc_style.add_style(
        'UpperLeftStyle', WD_STYLE_TYPE.CHARACTER)
    para_font = para_upper_left_style.font
    para_font.size = Pt(12)
    para_font.name = 'Arial'

    para_upper_left_style_spec = doc_style.add_style(
        'UpperLeftStyleSpec', WD_STYLE_TYPE.CHARACTER)
    para_font_spec = para_upper_left_style_spec.font
    para_font_spec.size = Pt(12)
    para_font_spec.name = 'Arial'

    heading_style = doc_style.add_style(
        'HeadingStyle', WD_STYLE_TYPE.CHARACTER)
    doc_font1 = heading_style.font
    doc_font1.size = Pt(14)
    doc_font1.name = 'Arial'
    doc_font1.bold = True

    normal_para_style = doc_style.add_style(
        'ParagraphStyle', WD_STYLE_TYPE.CHARACTER)
    normal_para_font = normal_para_style.font
    normal_para_font.size = Pt(11)
    normal_para_font.name = 'Arial'

    section = document.sections[0]
    header = section.header
    h = header.add_paragraph()
    h.add_run("AVV. ANTONIO MAZZA\n", "UpperLeftStyleSpec")
    h.add_run("www.legalars.net\nVia Lovanio n.10 – 20121 Milano\ntel.02.8295.1861 – mob.335.6260.014 – fax 178.6060.284\nemail: antonio.mazza@legalars.net – PEC: avv.antoniomazza@pec.giuffre.it", 'UpperLeftStyle')
    j = header.add_paragraph()

    paragraph_1 = document.add_paragraph()
    paragraph_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_1.paragraph_format.line_spacing = Pt(14)
    paragraph_1.add_run('ATTO DI PRECETTO', 'HeadingStyle')

    paragraph_2 = document.add_paragraph()
    paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_2.paragraph_format.line_spacing = Pt(22)
    if model.cr_tipo == 'Persona Fisica':
        paragraph_2.add_run("Il sottoscritto, Avv. Antonio MAZZA, nato a Catanzaro, il 6 febbraio 1967, codice fiscale MZZNTN67B06C352G, "
                            "con studio in Milano, Via Lovanio n.10, pec avv.antoniomazza@pec.giuffrè.it, in qualità di procuratore antistatario, "
                            f"nel procedimento monitorio n.PrAn.1 R.G. del CA di AA, di {model.cr_cognome} {model.cr_nome}, luogo di nascita {model.cr_luogo_di_nascita}, data nascita {model.cr_data_di_nascita}, "
                            f"residenza in {model.cr_comune_di_residenza}, {model.cr_indirizzo_di_residenza}, codice fiscale {model.cr_codice_fiscale}, rappresentato e difeso da sé medesimo, ex art.86 c.p.c., "
                            "ed elettivamente domiciliato presso il proprio studio in Milano – Via Lovanio 10;", 'ParagraphStyle')
    elif model.cr_tipo == 'Persona Giuridica':
        if model.cr_partita_iva is None:
            paragraph_2.add_run("Il sottoscritto, Avv. Antonio MAZZA, nato a Catanzaro, il 6 febbraio 1967, codice fiscale MZZNTN67B06C352G, "
                                "con studio in Milano, Via Lovanio n.10, pec avv.antoniomazza@pec.giuffrè.it, in qualità di procuratore antistatario, "
                                f"nel procedimento monitorio n.PrAn.1 R.G. del CA di AA, {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, "
                                f"codice fiscale {model.cr_codice_fiscale}, rappresentato e difeso da sé medesimo, ex art.86 c.p.c., ed elettivamente domiciliato presso "
                                "il proprio studio in Milano – Via Lovanio 10;", 'ParagraphStyle')
        elif model.cr_codice_fiscale is None:
            paragraph_2.add_run("Il sottoscritto, Avv. Antonio MAZZA, nato a Catanzaro, il 6 febbraio 1967, codice fiscale MZZNTN67B06C352G, "
                                "con studio in Milano, Via Lovanio n.10, pec avv.antoniomazza@pec.giuffrè.it, in qualità di procuratore antistatario, "
                                f"nel procedimento monitorio n.PrAn.1 R.G. del CA di AA, {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, "
                                f"partita IVA {model.cr_partita_iva}, rappresentato e difeso da sé medesimo, ex art.86 c.p.c., ed elettivamente domiciliato presso "
                                "il proprio studio in Milano – Via Lovanio 10;", 'ParagraphStyle')
        elif model.cr_partita_iva and model.cr_codice_fiscale is not None:
            paragraph_2.add_run("Il sottoscritto, Avv. Antonio MAZZA, nato a Catanzaro, il 6 febbraio 1967, codice fiscale MZZNTN67B06C352G, "
                                "con studio in Milano, Via Lovanio n.10, pec avv.antoniomazza@pec.giuffrè.it, in qualità di procuratore antistatario, "
                                f"nel procedimento monitorio n.PrAn.1 R.G. del CA di AA, {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, "
                                f"codice fiscale {model.cr_codice_fiscale}, partita IVA {model.cr_partita_iva}, rappresentato e difeso da sé medesimo, ex art.86 c.p.c., ed elettivamente domiciliato presso "
                                "il proprio studio in Milano – Via Lovanio 10;", 'ParagraphStyle')

    paragraph_3 = document.add_paragraph()
    paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_3.add_run('PREMESSO', 'HeadingStyle')
    paragraph_3.paragraph_format.line_spacing = Pt(24)

    paragraph_4 = document.add_paragraph()
    paragraph_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_4.paragraph_format.line_spacing = Pt(22)
    if model.cr_tipo == 'Persona Fisica':
        paragraph_4.add_run("il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, "
                            f"col quale il CA di AA, su ricorso presentato da {model.cr_cognome} {model.cr_nome}, luogo di nascita {model.cr_luogo_di_nascita},  data nascita {model.cr_data_di_nascita}, residenza in A.1.5, A.1.6, "
                            f"codice fiscale {model.cr_codice_fiscale}, ingiungeva a", 'ParagraphStyle')
    elif model.cr_tipo == 'Persona Giuridica':
        if model.cr_partita_iva is None:
            paragraph_4.add_run("il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, "
                                f"col quale il CA di AA, su ricorso presentato da {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, "
                                f"con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, ingiungeva a", 'ParagraphStyle')
        elif model.cr_codice_fiscale is None:
            paragraph_4.add_run("il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, "
                                f"col quale il CA di AA, su ricorso presentato da {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, "
                                f"con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, partita IVA {model.cr_partita_iva}, ingiungeva a", 'ParagraphStyle')
        elif model.cr_partita_iva and model.cr_codice_fiscale is not None:
            paragraph_4.add_run("il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, "
                                f"col quale il CA di AA, su ricorso presentato da {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, "
                                f"con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, partita IVA {model.cr_partita_iva}, ingiungeva a", 'ParagraphStyle')
    paragraph_5 = document.add_paragraph()
    paragraph_5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_5.paragraph_format.line_spacing = Pt(22)
    if model.db_pf or model.db_pj:
        if model.db_pf:
            for key, value in model.db_pf.items():
                print(value[f'db_nome'])
                paragraph_5.add_run(f"{value[f'db_cognome']} {value[f'db_nome']}, luogo nascita {value[f'db_luogo_di_nascita']}, data nascita {value[f'db_data_di_nascita']}, residenza in B.1.5, {value[f'db_indirizzo_di_residenza']}, "
                                    f"codice fiscale {value[f'df_codice_fiscale']}, ", 'ParagraphStyle')
        if model.db_pj:
            for key, value in model.db_pj.items():
                paragraph_5.add_run(
                    f"{value[f'db_denominazione_sociale']}, con sede in B.2.2, {value[f'db_sede_principale']}, codice fiscale {value[f'dj_codice_fiscale']}, partita IVA {value[f'dj_partita_iva']}, ", 'ParagraphStyle')

    paragraph_6 = document.add_paragraph()
    paragraph_6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_6.paragraph_format.line_spacing = Pt(22)
    paragraph_6.add_run(f"di pagare a parte ricorrente, per la causali di cui al ricorso, la somma di €. {model.somma}, oltre interessi Pr.4, "
                        f"spese vive della stessa procedura monitoria (liquidate in €. Pr.5) e spese e competenze successive occorrende, con ingiunzione di pagare pure "
                        f"il compenso per la procedura monitoria, liquidato in €. Pr.6, maggiorato del 15% per spese generali, dell’IVA e della CPA previsti per legge, "
                        f"e distratto ex art.93 c.p.c., al procuratore e difensore per non avere riscosso gli onorari; ", 'ParagraphStyle')

    paragraph_7 = document.add_paragraph()
    paragraph_7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_7.paragraph_format.line_spacing = Pt(24)
    paragraph_7.add_run('INTIMA', 'HeadingStyle')

    paragraph_8 = document.add_paragraph()
    paragraph_8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_8.paragraph_format.line_spacing = Pt(22)
    if model.db_pf or model.db_pj:
        if model.db_pf:
            for key, value in model.db_pf.items():
                paragraph_8.add_run(f"{value[f'db_cognome']} {value[f'db_nome']}, luogo nascita {value[f'db_luogo_di_nascita']}, data nascita {value[f'db_data_di_nascita']}, residenza in B.1.5, {value[f'db_indirizzo_di_residenza']}, "
                                    f"codice fiscale {value[f'df_codice_fiscale']}, ", 'ParagraphStyle')
        if model.db_pj:
            for key, value in model.db_pj.items():
                paragraph_8.add_run(
                    f"{value[f'db_denominazione_sociale']}, con sede in B.2.2, {value[f'db_sede_principale']}, codice fiscale {value[f'dj_codice_fiscale']}, partita IVA {value[f'dj_partita_iva']}, ", 'ParagraphStyle')

    paragraph_9 = document.add_table(rows=8, cols=2)
    row_0 = paragraph_9.rows[0]
    row_0.cells[0].add_paragraph(f'Compenso liquidato nel decreto ingiuntivo')
    zero = row_0.cells[1].add_paragraph(f'€. Pr.6')
    zero.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_1 = paragraph_9.rows[1]
    row_1.cells[0].add_paragraph(f'Spese generali (15% sul compenso)')
    one = row_1.cells[1].add_paragraph(f'PrAn.2')
    one.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_2 = paragraph_9.rows[2]
    row_2.cells[0].add_paragraph(
        f'CPA (4% sul totale compenso e spese generali)')
    two = row_2.cells[1].add_paragraph(f'PrAn.3')
    two.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_3 = paragraph_9.rows[3]
    row_3.cells[0].add_paragraph(
        f'Interessi legali sulle suindicate voci alla data Pr.7')
    three = row_3.cells[1].add_paragraph(f'PrAn. 4')
    three.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_4 = paragraph_9.rows[4]
    row_4.cells[0].add_paragraph(
        f'Spese vive successive alla emissione del decreto ingiuntivo e funzionali al presente precetto')
    four = row_4.cells[1].add_paragraph(f'PrAn. 5')
    four.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_5 = paragraph_9.rows[5]
    row_5.cells[0].add_paragraph(f'Compenso per il precetto')
    five = row_5.cells[1].add_paragraph(f'PrAn.6')
    five.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_6 = paragraph_9.rows[6]
    row_6.cells[0].add_paragraph(f'Spese generali (15% sul compenso precetto)')
    six = row_6.cells[1].add_paragraph(f'PrAn.7')
    six.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_7 = paragraph_9.rows[7]
    row_7.cells[0].add_paragraph(
        f'CPA (4% sul totale compenso precetto e spese generali)')
    seven = row_7.cells[1].add_paragraph(f'PrAn.8')
    seven.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph_10 = document.add_paragraph()
    paragraph_10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_10.paragraph_format.line_spacing = Pt(22)
    paragraph_10.add_run(f"e così complessivamente la somma di € PrAN.9 oltre agli interessi legai maturati successivamente "
                         "al Pr.7 fino all’effettivo e integrale soddisfo, nonché alle spese e competenze successive occorrende.", 'ParagraphStyle')

    paragraph_11 = document.add_paragraph()
    paragraph_11.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_11.paragraph_format.line_spacing = Pt(22)
    paragraph_11.add_run(f"Si avverte altresì parte debitrice che può, con l’ausilio di un organismo di composizione della crisi o di un "
                         f"professionista nominato dal giudice, porre rimedio alla situazione di sovraindebitamento, concludendo con la parte creditrice un accordo di "
                         f"composizione della crisi o proponendo alla medesima un piano del consumatore.", 'ParagraphStyle')

    paragraph_12 = document.add_paragraph()
    paragraph_12.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_12.paragraph_format.line_spacing = Pt(22)
    paragraph_12.add_run(
        f"Con espresso avvertimento che, in difetto, si procederà ad esecuzione forzata nei suoi confronti.", 'ParagraphStyle')

    paragraph_13 = document.add_paragraph()
    paragraph_13.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_13.paragraph_format.line_spacing = Pt(22)
    paragraph_13.add_run(f"Milano, Pr.7.", 'ParagraphStyle')

    paragraph_14 = document.add_paragraph()
    paragraph_14.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_14.paragraph_format.line_spacing = Pt(24)
    paragraph_14.add_run('(AVV. ANTONIO MAZZA)', 'HeadingStyle')

    # Save the document to a BytesIO object
    docx_file = io.BytesIO()
    document.save(docx_file)
    docx_file.seek(0)

    # Create a filename for the document
    filename = 'modello_atto_di_precetto_procuratore_antistatario.docx'

    response = FileResponse(docx_file, as_attachment=True, filename=filename)
    return response
