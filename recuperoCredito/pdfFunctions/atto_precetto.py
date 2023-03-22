import io
from django.http import FileResponse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def atto_precetto(model=None):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(4.75)
        section.bottom_margin = Cm(3.25)
        section.left_margin = Cm(2.75)
        section.right_margin = Cm(3.75)

    doc_style = document.styles

    para_upper_left_style = doc_style.add_style(
        'UpperLeftStyle', WD_STYLE_TYPE.CHARACTER)
    para_font = para_upper_left_style.font
    para_font.size = Pt(12)
    para_font.name = 'Arial'

    para_upper_left_style_spec = doc_style.add_style(
        'UpperLeftStyleSpec', WD_STYLE_TYPE.CHARACTER)
    para_font_spec = para_upper_left_style_spec.font
    para_font_spec.size = Pt(12)
    para_font_spec.name = 'Arial'

    heading_style = doc_style.add_style(
        'HeadingStyle', WD_STYLE_TYPE.CHARACTER)
    doc_font1 = heading_style.font
    doc_font1.size = Pt(14)
    doc_font1.name = 'Arial'
    doc_font1.bold = True

    normal_para_style = doc_style.add_style(
        'ParagraphStyle', WD_STYLE_TYPE.CHARACTER)
    normal_para_font = normal_para_style.font
    normal_para_font.size = Pt(11)
    normal_para_font.name = 'Arial'

    section = document.sections[0]
    header = section.header
    h = header.add_paragraph()
    h.add_run("AVV. ANTONIO MAZZA\n", "UpperLeftStyleSpec")
    h.add_run("www.legalars.net\nVia Lovanio n.10 – 20121 Milano\ntel.02.8295.1861 – mob.335.6260.014 – fax 178.6060.284\nemail: antonio.mazza@legalars.net – PEC: avv.antoniomazza@pec.giuffre.it", 'UpperLeftStyle')
    j = header.add_paragraph()

    paragraph_1 = document.add_paragraph()
    paragraph_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_1.paragraph_format.line_spacing = Pt(14)
    paragraph_1.add_run('ATTO DI PRECETTO', 'HeadingStyle')

    paragraph_2 = document.add_paragraph()
    paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_2.paragraph_format.line_spacing = Pt(22)
    if model.cr_tipo == 'Persona Fisica':
        paragraph_2.add_run(f"{model.cr_cognome} {model.cr_nome}, luogo nascita {model.cr_luogo_di_nascita}, data nascita {model.cr_data_di_nascita}, residenza in {model.cr_comune_di_residenza}, {model.cr_indirizzo_di_residenza}, codice fiscale {model.cr_codice_fiscale}, ", 'ParagraphStyle')
    elif model.cr_tipo == 'Persona Giuridica':
        if model.cr_partita_iva is None:
            paragraph_2.add_run(
                f"{model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, ", 'ParagraphStyle')
        elif model.cr_codice_fiscale is None:
            paragraph_2.add_run(
                f"{model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, partita IVA {model.cr_partita_iva}, ", 'ParagraphStyle')
        elif model.cr_partita_iva and model.cr_codice_fiscale is not None:
            paragraph_2.add_run(f"{model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, partita IVA {model.cr_partita_iva}, ", 'ParagraphStyle')

    paragraph_3 = document.add_paragraph()
    paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_3.paragraph_format.line_spacing = Pt(22)
    paragraph_3.add_run("con domicilio eletto in Milano – Via Lovanio 10, presso lo studio dell’Avv. Antonio MAZZA (codice fiscale MZZNTN67B06C352G, PEC avv.antoniomazza@pec.giuffre.it, fax 1786060284), che ha specifico mandato di rappresentanza e difesa, giusta procura speciale allegata al ricorso volto a ottenere il decreto ingiuntivo qui di seguito richiamato; ", 'ParagraphStyle')

    paragraph_4 = document.add_paragraph()
    paragraph_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_4.paragraph_format.line_spacing = Pt(14)
    paragraph_4.add_run('PREMESSO', 'HeadingStyle')

    paragraph_5 = document.add_paragraph()
    paragraph_5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_5.paragraph_format.line_spacing = Pt(22)
    if model.cr_tipo == 'Persona Fisica':
        paragraph_5.add_run(
            f"il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, col quale il CA di AA, su ricorso presentato da {model.cr_cognome} {model.cr_nome}, luogo nascita {model.cr_luogo_di_nascita}, data nascita {model.cr_data_di_nascita}, residenza in {model.cr_comune_di_residenza}, {model.cr_indirizzo_di_residenza}, codice fiscale {model.cr_codice_fiscale}, ingiungeva a ", 'ParagraphStyle')
    elif model.cr_tipo == 'Persona Giuridica':
        if model.cr_partita_iva is None:
            paragraph_5.add_run(
                f"il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, col quale il CA di AA, su ricorso presentato da {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, ingiungeva a ", 'ParagraphStyle')
        elif model.cr_codice_fiscale is None:
            paragraph_5.add_run(
                f"il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, col quale il CA di AA, su ricorso presentato da {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, partita IVA {model.cr_partita_iva}, ingiungeva a ", 'ParagraphStyle')
        elif model.cr_partita_iva and model.cr_codice_fiscale is not None:
            paragraph_5.add_run(
                f"il contenuto -da intendersi qui integralmente trascritto- del decreto ingiuntivo n. Pr.1, notificato il Pr.2, esecutivo con apposizione della formula di rito in data Pr.3, col quale il CA di AA, su ricorso presentato da {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, partita IVA {model.cr_partita_iva}, ingiungeva a ", 'ParagraphStyle')

    paragraph_6 = document.add_paragraph()
    paragraph_6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_6.paragraph_format.line_spacing = Pt(22)
    if model.db_pf or model.db_pj:
        if model.db_pf:
            for key, value in model.db_pf.items():
                paragraph_6.add_run(f"{value[f'db_cognome']} {value[f'db_nome']}, luogo nascita {value[f'db_luogo_di_nascita']}, data nascita {value[f'db_data_di_nascita']}, residenza in B.1.5, {value[f'db_indirizzo_di_residenza']}, "
                                    f"codice fiscale {value[f'df_codice_fiscale']}, ", 'ParagraphStyle')
        if model.db_pj:
            for key, value in model.db_pj.items():
                paragraph_6.add_run(
                    f"{value[f'db_denominazione_sociale']}, con sede in B.2.2, {value[f'db_sede_principale']}, codice fiscale {value[f'dj_codice_fiscale']}, partita IVA {value[f'dj_partita_iva']}, ", 'ParagraphStyle')

    paragraph_7 = document.add_paragraph()
    paragraph_7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_7.paragraph_format.line_spacing = Pt(22)
    paragraph_7.add_run(f"di pagare a parte ricorrente per la causali di cui al ricorso la somma di €. C, oltre interessi Pr.4, spese vive della stessa procedura monitoria (liquidate in €. Pr.5) e spese e competenze successive occorrende, con ingiunzione di pagare pure il compenso per la procedura monitoria, liquidato in €. Pr.6, maggiorato del 15% per spese generali, dell’IVA e della CPA previsti per legge, e distratto ex art.93 c.p.c., al procuratore e difensore per non avere riscosso gli onorari;", 'ParagraphStyle')

    paragraph_8 = document.add_paragraph()
    paragraph_8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_8.paragraph_format.line_spacing = Pt(14)
    paragraph_8.add_run('INTIMA', 'HeadingStyle')

    paragraph_9 = document.add_paragraph()
    paragraph_9.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_9.paragraph_format.line_spacing = Pt(22)
    if model.db_pf or model.db_pj:
        if model.db_pf:
            for key, value in model.db_pf.items():
                paragraph_9.add_run(f"a {value[f'db_cognome']} {value[f'db_nome']}, luogo nascita {value[f'db_luogo_di_nascita']}, data nascita {value[f'db_data_di_nascita']}, residenza in B.1.5, {value[f'db_indirizzo_di_residenza']}, "
                                    f"codice fiscale {value[f'df_codice_fiscale']}, ", 'ParagraphStyle')
        if model.db_pj:
            for key, value in model.db_pj.items():
                paragraph_9.add_run(
                    f"a {value[f'db_denominazione_sociale']}, con sede in B.2.2, {value[f'db_sede_principale']}, codice fiscale {value[f'dj_codice_fiscale']}, partita IVA {value[f'dj_partita_iva']}, ", 'ParagraphStyle')

    paragraph_10 = document.add_paragraph()
    paragraph_10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_10.paragraph_format.line_spacing = Pt(22)
    if model.cr_tipo == 'Persona Fisica':
        paragraph_10.add_run(
            f"il pagamento, in favore di {model.cr_cognome} {model.cr_nome}, luogo nascita {model.cr_luogo_di_nascita}, data nascita {model.cr_data_di_nascita}, residenza in {model.cr_comune_di_residenza}, {model.cr_indirizzo_di_residenza}, codice fiscale {model.cr_codice_fiscale}, entro dieci giorni dalla notifica del presente atto, delle seguenti somme: ", 'ParagraphStyle')
    elif model.cr_tipo == 'Persona Giuridica':
        if model.cr_partita_iva is None:
            paragraph_10.add_run(
                f"il pagamento, in favore di {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, entro dieci giorni dalla notifica del presente atto, delle seguenti somme: ", 'ParagraphStyle')
        elif model.cr_codice_fiscale is None:
            paragraph_10.add_run(
                f"il pagamento, in favore di {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, partita IVA {model.cr_partita_iva}, entro dieci giorni dalla notifica del presente atto, delle seguenti somme: ", 'ParagraphStyle')
        elif model.cr_partita_iva and model.cr_codice_fiscale is not None:
            paragraph_10.add_run(
                f"il pagamento, in favore di {model.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {model.cr_comune_sede_principale}, {model.cr_indirizzo_sede_principale}, codice fiscale {model.cr_codice_fiscale}, partita IVA {model.cr_partita_iva}, entro dieci giorni dalla notifica del presente atto, delle seguenti somme: ", 'ParagraphStyle')

    paragraph_11 = document.add_table(rows=7, cols=2)
    row_0 = paragraph_11.rows[0]
    row_0.cells[0].add_paragraph(f'Capitale ')
    zero = row_0.cells[1].add_paragraph(f'€. {model.somma}')
    zero.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_1 = paragraph_11.rows[1]
    row_1.cells[0].add_paragraph(f'Interessi Pr.4 alla data Pr.7')
    one = row_1.cells[1].add_paragraph(f'Pr. 8')
    one.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_2 = paragraph_11.rows[2]
    row_2.cells[0].add_paragraph(
        f"Spese vive liquidate nell'ingiunzione ")
    two = row_2.cells[1].add_paragraph(f' Pr.5')
    two.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_3 = paragraph_11.rows[3]
    row_3.cells[0].add_paragraph(
        f'Spese vive successive alla emissione del decreto ingiuntivo')
    three = row_3.cells[1].add_paragraph(f'Pr.9')
    three.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_4 = paragraph_11.rows[4]
    row_4.cells[0].add_paragraph(
        f'Compenso per il precetto')
    four = row_4.cells[1].add_paragraph(f'Pr.10')
    four.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_5 = paragraph_11.rows[5]
    row_5.cells[0].add_paragraph(f'Spese generali (15% sul compenso precetto)')
    five = row_5.cells[1].add_paragraph(f'Pr.11')
    five.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    row_6 = paragraph_11.rows[6]
    row_6.cells[0].add_paragraph(
        f'CPA (4% sul totale compenso precetto e spese generali)')
    six = row_6.cells[1].add_paragraph(f'Pr.12')
    six.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph_12 = document.add_
