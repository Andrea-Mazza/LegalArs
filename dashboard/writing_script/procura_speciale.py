import io
from django.http import FileResponse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from django.views import View
import os
import uuid


def write_procura_speciale(pratica):
    # num will be concatenated to filename at the end of this script

    # Create a new Word document
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(4.75)
        section.bottom_margin = Cm(3.25)
        section.left_margin = Cm(2.75)
        section.right_margin = Cm(3.75)
    doc_style = document.styles
    para_upper_left_style = doc_style.add_style(
        'UpperLeftStyle', WD_STYLE_TYPE.CHARACTER)
    para_font = para_upper_left_style.font
    para_font.size = Pt(12)
    para_font.name = 'Arial'
    para_upper_left_style_spec = doc_style.add_style(
        'UpperLeftStyleSpec', WD_STYLE_TYPE.CHARACTER)
    para_font_spec = para_upper_left_style_spec.font
    para_font_spec.size = Pt(12)
    para_font_spec.name = 'Arial'
    heading_style = doc_style.add_style(
        'HeadingStyle', WD_STYLE_TYPE.CHARACTER)
    doc_font1 = heading_style.font
    doc_font1.size = Pt(14)
    doc_font1.name = 'Arial'
    doc_font1.bold = True
    normal_para_style = doc_style.add_style(
        'ParagraphStyle', WD_STYLE_TYPE.CHARACTER)
    normal_para_font = normal_para_style.font
    normal_para_font.size = Pt(11)
    normal_para_font.name = 'Arial'
    section = document.sections[0]
    header = section.header
    h = header.add_paragraph()
    h.add_run("AVV. ANTONIO MAZZA\n", "UpperLeftStyleSpec")
    h.add_run("www.legalars.net\nVia Lovanio n.10 – 20121 Milano\ntel.02.8295.1861 – mob.335.6260.014 – fax 178.6060.284\nemail: antonio.mazza@legalars.net – PEC: avv.antoniomazza@pec.giuffre.it", 'UpperLeftStyle')
    j = header.add_paragraph()

    paragraph_1 = document.add_paragraph()
    paragraph_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_1.paragraph_format.line_spacing = Pt(14)
    paragraph_1.add_run('PROCURA SPECIALE', 'HeadingStyle')
    # paragraph_1.styles.style.font.name = 'Arial'
    # paragraph_1.styles.style.font.size = Pt(14)
    # paragraph_1.add_run("ATTO DI PRECETTO").bold = True
    paragraph_2 = document.add_paragraph()
    paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_2.paragraph_format.line_spacing = Pt(22)
    if pratica.cr_tipo == 'Persona Fisica':
        if not pratica.cr_partita_iva:
            paragraph_2.add_run(
                f"{pratica.cr_nome} {pratica.cr_cognome}, {pratica.cr_luogo_di_nascita}, {pratica.cr_data_di_nascita}, {pratica.cr_comune_di_residenza}, {pratica.cr_indirizzo_di_residenza}, {pratica.cr_codice_fiscale}, ", 'ParagraphStyle')
        elif not pratica.cr_codice_fiscale:
            paragraph_2.add_run(
                f"{pratica.cr_nome} {pratica.cr_cognome}, {pratica.cr_luogo_di_nascita}, {pratica.cr_data_di_nascita}, {pratica.cr_comune_di_residenza}, {pratica.cr_indirizzo_di_residenza}, {pratica.cr_partita_iva}, ", 'ParagraphStyle')
        elif pratica.cr_codice_fiscale and pratica.cr_partita_iva:
            paragraph_2.add_run(f"{pratica.cr_nome} {pratica.cr_cognome}, {pratica.cr_luogo_di_nascita}, {pratica.cr_data_di_nascita}, {pratica.cr_comune_di_residenza}, {pratica.cr_indirizzo_di_residenza}, {pratica.cr_codice_fiscale}, {pratica.cr_partita_iva}, ", 'ParagraphStyle')
    elif pratica.cr_tipo == 'Persona Giuridica':
        if not pratica.cr_partita_iva:
            paragraph_2.add_run(
                f"{pratica.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {pratica.cr_comune_sede_principale}, {pratica.cr_indirizzo_sede_principale}, {pratica.cr_codice_fiscale},", 'ParagraphStyle')
        elif not pratica.cr_codice_fiscale:
            paragraph_2.add_run(
                f"{pratica.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {pratica.cr_comune_sede_principale}, {pratica.cr_indirizzo_sede_principale}, {pratica.cr_partita_iva},", 'ParagraphStyle')
        elif pratica.cr_partita_iva and pratica.cr_codice_fiscale:
            paragraph_2.add_run(f"{pratica.cr_denominazione_sociale}, in persona del legale rappresentante pro tempore, con sede in {pratica.cr_comune_sede_principale}, {pratica.cr_indirizzo_sede_principale}, {pratica.cr_codice_fiscale}, {pratica.cr_partita_iva}", 'ParagraphStyle')

    paragraph_8 = document.add_paragraph()
    paragraph_8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_8.paragraph_format.line_spacing = Pt(22)
    if pratica.db_tipo == 'Persona Fisica':
        paragraph_8.add_run(f"nomina proprio procuratore e difensore, per l’attività stragiudiziale e nel procedimento monitorio, volti al recupero di credito nei confronti di {pratica.db_cognome} {pratica.db_nome}, luogo nascita {pratica.db_luogo_di_nascita}, data nascita {pratica.db_data_di_nascita}, residenza in {pratica.db_comune_di_residenza}, {pratica.db_indirizzo_di_residenza}, "
                            f"codice fiscale {pratica.df_codice_fiscale}, ", 'ParagraphStyle')
    else:
        paragraph_8.add_run(
            f"nomina proprio procuratore e difensore, per l’attività stragiudiziale e nel procedimento monitorio, volti al recupero di credito nei confronti di {pratica.db_denominazione_sociale}, con sede in {pratica.db_comune_sede_principale}, {pratica.db_indirizzo_sede_principale}, codice fiscale {pratica.dj_codice_fiscale}, partita IVA {pratica.dj_partita_iva}, ", 'ParagraphStyle')

    paragraph_9 = document.add_paragraph()
    paragraph_9.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    paragraph_9.paragraph_format.line_spacing = Pt(22)
    paragraph_9.add_run(f"l'Avv. Antonio MAZZA, con studio in Milano, Via Lovanio n.10, con ogni più ampio potere, incluso quello di transigere e conciliare, riscuotere e quietanzare, rinunciare agli atti e farsi sostituire. Il mandante elegge domicilio presso il predetto studio legale di Milano, Via Lovanio n.10, nonché all’indirizzo avv.antoniomazza@pec.giuffre.it, per le comunicazioni via PEC. Con la sottoscrizione del presente atto sI dichiara di aver preso visione dell'informativa resa ai sensi del D. Lgs. 30 giungo 2003, n. 196 art. 13 e si autorizza il trattamento dei relativi dati per le finalità di cui al presente mandato. Si allega copia della carta di identità del sottoscrittore.", 'ParagraphStyle')

    paragraph_10 = document.add_paragraph()
    paragraph_10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_10.paragraph_format.line_spacing = Pt(22)
    paragraph_10.paragraph_format.space_after = Pt(42)
    paragraph_10.add_run('Firma per esteso del mandante', 'HeadingStyle')

    paragraph_11 = document.add_paragraph()
    paragraph_11.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_11.paragraph_format.line_spacing = Pt(22)
    paragraph_11.add_run('Visto per autentica della firma', 'HeadingStyle')

    paragraph_14 = document.add_paragraph()
    paragraph_14.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_14.paragraph_format.line_spacing = Pt(24)
    paragraph_14.add_run('(AVV. ANTONIO MAZZA)', 'HeadingStyle')
    # Save the document to a BytesIO object
    docx_file = io.BytesIO()
    document.save(docx_file)
    docx_file.seek(0)
    # Create a filename for the document
    # unique_id = str(uuid.uuid4())
    # filename = f"procura_speciale_{unique_id}.docx"
    # # response = FileResponse(docx_file, as_attachment=True, filename=filename)
    # # return response
    # # Crea il percorso completo del file
    # file_path = os.path.join(os.path.dirname(os.path.abspath(
    #     __file__)), '..', 'recupero_credito_doc', filename)

    # # Salva il file sul disco
    # with open(file_path, 'wb') as f:
    #     f.write(docx_file.getbuffer())
    return docx_file
